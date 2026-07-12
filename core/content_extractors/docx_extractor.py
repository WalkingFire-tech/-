"""
Word文档提取器 - 支持doc和docx
"""

from pathlib import Path
from typing import Optional, List
from loguru import logger

from .base import ContentExtractor


class DocxExtractor(ContentExtractor):
    """Word文档提取器"""
    
    def __init__(self):
        self._available = False
        self._docx = None
        
        try:
            import docx
            self._docx = docx
            self._available = True
            logger.info("python-docx 可用，Word文档提取已启用")
        except ImportError:
            logger.debug("python-docx 未安装，DOCX 提取不可用")
    
    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ('.docx', '.doc') and self._available
    
    def extract(self, file_path: Path) -> Optional[str]:
        if not self._available:
            return None
        
        try:
            doc = self._docx.Document(str(file_path))
            
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts) if text_parts else ''
            
        except Exception as e:
            logger.error(f"DOCX 提取失败 {file_path}: {e}")
            return None
    
    def get_supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']
    
    def get_name(self) -> str:
        return "DocxExtractor"